from pathlib import Path

import pytest
from docx import Document as DocxDocument

from scrapenfill.core.process import Process


def test_extract_text_from_txt(config, txt_file: Path):
    process = Process(config)
    result = process.extract_text(txt_file)
    assert "Jan" in result
    assert "Jansen" in result
    assert "Senior Software Engineer" in result


def test_extract_text_from_docx(config, docx_file: Path):
    process = Process(config)
    result = process.extract_text(docx_file)
    assert "Jan" in result
    assert "Jansen" in result


def test_extract_text_unknown_extension(config, tmp_path: Path):
    path = tmp_path / "cv.unknown"
    path.write_text("test")
    process = Process(config)
    result = process.extract_text(path)
    assert result == ""


def test_save_docx(config, docx_template: Path, tmp_path: Path):
    data = {"voornaam": "Jan", "achternaam": "Jansen", "functie_titel": "Engineer"}
    output_path = tmp_path / "output.docx"

    process = Process(config)
    process.save_docx(data, str(docx_template), output_path)

    assert output_path.exists()
    doc = DocxDocument(str(output_path))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "Jan" in text
    assert "Jansen" in text
    assert "Engineer" in text


def test_save_docx_overwrites_existing(config, docx_template: Path, tmp_path: Path):
    data = {"voornaam": "Piet", "achternaam": "Pieters", "functie_titel": "Dev"}
    output_path = tmp_path / "output.docx"
    output_path.write_text("stale")

    process = Process(config)
    process.save_docx(data, str(docx_template), output_path)

    assert output_path.exists()
    doc = DocxDocument(str(output_path))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "Piet" in text


@pytest.mark.parametrize(
    "provider,expected_client",
    [
        ("CHATGPT", "OpenAIClient"),
        ("MISTRAL", "MistralClient"),
        ("GEMINI", "GeminiClient"),
        ("OLLAMA", "OllamaClient"),
        ("CLAUDE", "ClaudeClient"),
    ],
)
async def test_cv_to_json_dispatches_correct_client(
    config, provider, expected_client, model_schema, prompt_file, tmp_path, monkeypatch
):
    config["LLM"]["provider"] = provider

    core_dir = tmp_path / "core"
    core_dir.mkdir()
    core_dir.joinpath("model").write_text(
        '{"type":"json_schema","json_schema":{"name":"x","schema":{"type":"object","properties":{"voornaam":{"type":"string"},"achternaam":{"type":"string"}},"required":["voornaam","achternaam"]}}}'
    )
    core_dir.joinpath("prompt").write_text("Extraheer CV data naar JSON.")

    called_with = {}

    class MockClient:
        async def extract(self, prompt, schema):
            called_with["prompt"] = prompt
            called_with["schema"] = schema
            return {"voornaam": "Jan", "achternaam": "Jansen"}

    monkeypatch.setattr(f"scrapenfill.core.process.{expected_client}", lambda config: MockClient())

    monkeypatch.chdir(tmp_path)
    process = Process(config)
    result = await process.cv_to_json("Jan Jansen CV text")

    assert result == {"voornaam": "Jan", "achternaam": "Jansen"}


async def test_cv_to_json_raises_on_invalid_provider(config, tmp_path, monkeypatch):
    config["LLM"]["provider"] = "INVALID"
    core_dir = tmp_path / "core"
    core_dir.mkdir()
    core_dir.joinpath("model").write_text(
        '{"type":"json_schema","json_schema":{"name":"x","schema":{"type":"object","properties":{}}}}'
    )
    core_dir.joinpath("prompt").write_text("test prompt")
    monkeypatch.chdir(tmp_path)

    process = Process(config)
    with pytest.raises(Exception, match="Invalid LLM provider"):
        await process.cv_to_json("some text")


async def test_cv_to_json_returns_none_on_empty_response(config, tmp_path, monkeypatch):
    config["LLM"]["provider"] = "CHATGPT"
    core_dir = tmp_path / "core"
    core_dir.mkdir()
    core_dir.joinpath("model").write_text(
        '{"type":"json_schema","json_schema":{"name":"x","schema":{"type":"object","properties":{}}}}'
    )
    core_dir.joinpath("prompt").write_text("test prompt")

    class MockClient:
        async def extract(self, prompt, schema):
            return {}

    monkeypatch.setattr("scrapenfill.core.process.OpenAIClient", lambda config: MockClient())
    monkeypatch.chdir(tmp_path)

    process = Process(config)
    result = await process.cv_to_json("some text")
    assert result == {}


async def test_cv_to_json_reads_model_and_prompt_from_cwd(config, tmp_path, monkeypatch):
    config["LLM"]["provider"] = "CHATGPT"
    core_dir = tmp_path / "core"
    core_dir.mkdir()
    core_dir.joinpath("model").write_text(
        '{"type":"json_schema","json_schema":{"name":"x","schema":{"type":"object","properties":{"voornaam":{"type":"string"}}}}}'
    )
    core_dir.joinpath("prompt").write_text("custom prompt")

    class MockClient:
        async def extract(self, prompt, schema):
            return {"voornaam": "Test"}

    monkeypatch.setattr("scrapenfill.core.process.OpenAIClient", lambda config: MockClient())
    monkeypatch.chdir(tmp_path)

    process = Process(config)
    result = await process.cv_to_json("input")
    assert result == {"voornaam": "Test"}
