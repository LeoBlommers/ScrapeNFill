import json
from configparser import ConfigParser
from pathlib import Path

import pytest
from docx import Document


@pytest.fixture
def config() -> ConfigParser:
    cfg = ConfigParser()
    cfg.read_string(
        """
[DIRECTORIES]
input = /tmp/input
output = /tmp/output

[TEMPLATE]
template = /tmp/template.docx

[LLM]
provider = CHATGPT

[CHATGPT]
model = gpt-5.4-mini
api_key = test-key

[MISTRAL]
model = mistral-large-latest
api_key = test-key

[GEMINI]
model = gemini-3.5-flash
api_key = test-key

[OLLAMA]
model = qwen3:32b
url = http://localhost:11434

[CLAUDE]
model = claude-sonnet-4-20250514
api_key = test-key

[PROCESSING]
max_concurrent = 5
"""
    )
    return cfg


@pytest.fixture
def cv_text() -> str:
    return """Voornaam: Jan
Achternaam: Jansen
Email: jan.jansen@email.com
Telefoon: 06-12345678
Functie: Senior Software Engineer

Samenvatting: Ervaren developer met 10 jaar ervaring in Python en JavaScript.

Werkervaring:
- Bedrijf: TechCorp (2020-2024)
  Rol: Senior Developer
  Beschrijving: Leidde een team van 5 developers.
  Skills: Python, Docker, AWS
"""


@pytest.fixture
def txt_file(tmp_path: Path, cv_text: str) -> Path:
    path = tmp_path / "cv.txt"
    path.write_text(cv_text, encoding="utf-8")
    return path


@pytest.fixture
def docx_file(tmp_path: Path, cv_text: str) -> Path:
    path = tmp_path / "cv.docx"
    doc = Document()
    for line in cv_text.strip().split("\n"):
        doc.add_paragraph(line.strip())
    doc.save(str(path))
    return path


@pytest.fixture
def docx_template(tmp_path: Path) -> Path:
    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{ voornaam }} {{ achternaam }}")
    doc.add_paragraph("{{ functie_titel }}")
    doc.save(str(path))
    return path


@pytest.fixture
def model_schema(tmp_path: Path) -> dict:
    schema = {
        "type": "json_schema",
        "json_schema": {
            "name": "cv_data",
            "schema": {
                "type": "object",
                "properties": {
                    "voornaam": {"type": "string"},
                    "achternaam": {"type": "string"},
                },
                "required": ["voornaam", "achternaam"],
            },
        },
    }
    path = tmp_path / "model"
    path.write_text(json.dumps(schema))
    return schema


@pytest.fixture
def prompt_file(tmp_path: Path) -> str:
    prompt = "Extraheer CV data naar JSON."
    path = tmp_path / "prompt"
    path.write_text(prompt)
    return prompt
